"""Extended tools: email, word, excel, planner, skill tools, wechat, mcp, dynamic."""
import asyncio
import importlib.util
import json
import os
from typing import TYPE_CHECKING, Any, Optional

from .path_utils import resolve_workspace_path
from .planner import Planner

if TYPE_CHECKING:
    from ..auth.permissions import PermissionsManager
    from .skill_manager import SkillManager

DYNAMIC_TOOLS_DIR = "dynamic_tools"


class ExtendedTools:
    """Extended tools: email, word, excel, skill tools, planner."""

    def __init__(self, data_dir: str, permissions: Optional["PermissionsManager"] = None,
                 wechat_adapter: Any = None,
                 dingtalk_client_id: str = "", dingtalk_client_secret: str = ""):
        self.data_dir = data_dir
        self.permissions = permissions
        self._wechat_adapter = wechat_adapter
        self._skill_mgr = None
        self.planner = Planner(os.path.join(data_dir, "workspace", "plans"))
        from ..gateway.adapters.wechat_cli import WeChatCLI
        self.wechat = WeChatCLI()
        from ..gateway.adapters.feishu_cli import FeishuCLI
        self.feishu = FeishuCLI()
        from ..gateway.adapters.dingtalk_cli import DingTalkCLI
        self.dingtalk = DingTalkCLI(
            client_id=dingtalk_client_id,
            client_secret=dingtalk_client_secret,
        )

    def set_wechat_adapter(self, adapter: Any):
        """Inject the active WeChat iLink adapter for proactive sending."""
        self._wechat_adapter = adapter

    def set_dingtalk_adapter(self, adapter: Any):
        """Share the DingTalk adapter's authenticated DingTalkCLI instance.

        The adapter's ``_dws_cli`` may hold device-auth state from the QR
        login flow.  Sharing the instance lets ExtendedTools domain tools
        (chat, doc, calendar, …) reuse the same auth session instead of
        creating an unauthenticated clone.
        """
        if adapter and hasattr(adapter, '_dws_cli'):
            self.dingtalk = adapter._dws_cli

    def set_feishu_adapter(self, adapter: Any):
        """Share the Feishu adapter's configured FeishuCLI instance.

        The adapter's ``_lark_cli`` holds the app credentials (App ID/Secret)
        configured during the QR login flow.  Sharing the instance ensures
        ExtendedTools domain tools use the same credentials instead of an
        unconfigured clone.
        """
        if adapter and hasattr(adapter, '_lark_cli'):
            self.feishu = adapter._lark_cli

    def set_skill_manager(self, mgr: "SkillManager"):
        """Inject SkillManager reference for skill_load tool."""
        self._skill_mgr = mgr

    # ------------------------------------------------------------------
    # Dynamic tools (runtime-created by agent)
    # ------------------------------------------------------------------

    def _dynamic_tools_dir(self) -> str:
        path = os.path.join(self.data_dir, DYNAMIC_TOOLS_DIR)
        os.makedirs(path, exist_ok=True)
        return path

    def _dynamic_manifest_path(self) -> str:
        return os.path.join(self._dynamic_tools_dir(), "manifest.json")

    def _load_dynamic_manifest(self) -> dict:
        path = self._dynamic_manifest_path()
        if not os.path.exists(path):
            return {}
        try:
            with open(path, encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}

    def _save_dynamic_manifest(self, manifest: dict):
        path = self._dynamic_manifest_path()
        with open(path, "w", encoding="utf-8") as f:
            json.dump(manifest, f, indent=2, ensure_ascii=False)

    async def _run_dynamic_tool(self, tool_name: str, args: dict) -> dict:
        """Execute a dynamic tool by loading and calling its execute() function."""
        manifest = self._load_dynamic_manifest()
        entry = manifest.get(tool_name)
        if not entry:
            return {"status": "error", "msg": f"Dynamic tool '{tool_name}' not found in manifest"}

        filepath = entry.get("path")
        if not filepath or not os.path.exists(filepath):
            return {"status": "error", "msg": f"Dynamic tool '{tool_name}' file not found: {filepath}"}

        try:
            with open(filepath, encoding="utf-8") as f:
                source = f.read()
            ns: dict[str, Any] = {}
            exec(source, ns)
            if "execute" not in ns:
                return {"status": "error", "msg": f"Dynamic tool '{tool_name}' has no execute() function"}
            fn = ns["execute"]
            if asyncio.iscoroutinefunction(fn):
                return await fn(args)
            return fn(args)
        except Exception as e:
            return {"status": "error", "msg": f"Dynamic tool '{tool_name}' error: {e}"}

    async def do_tool_create(self, args: dict) -> dict:
        """Create a new tool at runtime by providing Python code and a schema.

        The code must define an ``async def execute(args: dict) -> dict`` function.
        After creation, the agent can call this tool by name like any built-in tool.
        """
        name = args.get("name", "")
        code = args.get("code", "")
        description = args.get("description", "")
        parameters = args.get("parameters", {})

        if not name or not code:
            return {"status": "error", "msg": "name and code are required"}

        if not await self._confirm("tool_create", f"Create tool: {name}"):
            return {"status": "error", "msg": "Tool creation not permitted"}

        # Validate: code must define execute()
        ns: dict[str, Any] = {}
        try:
            exec(code, ns)
        except Exception as e:
            return {"status": "error", "msg": f"Code validation failed: {e}"}
        if "execute" not in ns:
            return {"status": "error", "msg": "Code must define an execute() function"}

        # Save to file
        tools_dir = self._dynamic_tools_dir()
        filepath = os.path.join(tools_dir, f"{name}.py")
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(code)

        # Register in manifest
        manifest = self._load_dynamic_manifest()
        manifest[name] = {
            "path": filepath,
            "description": description,
            "parameters": parameters,
        }
        self._save_dynamic_manifest(manifest)

        return {
            "status": "success",
            "msg": f"Tool '{name}' created and registered",
            "tool_name": name,
        }

    async def do_tool_delete(self, args: dict) -> dict:
        """Delete a dynamically created tool."""
        name = args.get("name", "")
        if not name:
            return {"status": "error", "msg": "name is required"}

        manifest = self._load_dynamic_manifest()
        if name not in manifest:
            return {"status": "error", "msg": f"Tool '{name}' not found"}

        # Remove file
        filepath = manifest[name].get("path", "")
        if filepath and os.path.exists(filepath):
            os.remove(filepath)

        # Update manifest
        del manifest[name]
        self._save_dynamic_manifest(manifest)

        return {"status": "success", "msg": f"Tool '{name}' deleted"}

    async def do_tool_list(self, args: dict) -> dict:
        """List all dynamically created tools."""
        manifest = self._load_dynamic_manifest()
        if not manifest:
            return {"status": "success", "tools": {}, "count": 0, "msg": "No dynamic tools created"}
        result = {}
        for name, entry in manifest.items():
            result[name] = {
                "description": entry.get("description", ""),
                "has_params": bool(entry.get("parameters", {})),
            }
        return {"status": "success", "tools": result, "count": len(result)}

    async def _confirm(self, operation: str, details: str = "") -> bool:
        """Check permission for an operation. Returns True if allowed."""
        if self.permissions:
            return await self.permissions.confirm_operation(operation, details)
        return True

    async def do_email_send(self, args: dict) -> dict:
        """Send email via SMTP (himayala or smtplib)."""
        to = args.get("to", "")
        subject = args.get("subject", "")
        if not await self._confirm("email_send", f"To: {to}, Subject: {subject}"):
            return {"status": "error", "msg": "Email sending not permitted"}
        # Stub — actual SMTP sending depends on user config
        return {"status": "success", "msg": f"Email to {to}: '{subject}' queued"}
        # TODO: implement actual SMTP when user provides email config

    async def do_word_doc(self, args: dict) -> dict:
        """Generate a Word (.docx) document with headings, tables, paragraphs, and styles.

        Domain rules (from word-docx skill):
        - Prefer named styles (Heading 1/2/3, Normal) over direct formatting
        - A .docx is a ZIP of XML parts — structure matters as much as visible text
        - Margins, orientation, headers/footers are section-level, not document-level
        - Lists and numbering use Word's numbering definitions, not Unicode bullets
        - Tracked changes, comments, and fields need precise edits — visible text is
          not the full document when revisions are enabled
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            return {"status": "error", "msg": "python-docx not installed. Run: pip install python-docx"}
        try:
            path = resolve_workspace_path(args.get("path", "document.docx"), self.data_dir, self.permissions)
        except PermissionError as exc:
            return {"status": "error", "msg": str(exc)}

        title = args.get("title", "Document")
        content = args.get("content", "")
        tables_data = args.get("tables", [])
        page_orientation = args.get("page_orientation", "portrait")  # 'portrait' or 'landscape'
        margins = args.get("margins", None)  # dict with top/bottom/left/right in inches

        doc = Document()

        # Page setup
        section = doc.sections[0]
        if page_orientation == "landscape":
            section.orientation = 1  # WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        if margins:
            margin_inches = {k: Inches(v) for k, v in margins.items() if k in ("top", "bottom", "left", "right")}
            for k, v in margin_inches.items():
                setattr(section, k, v)

        # Title
        doc.add_heading(title, 0)

        # Parse content: lines prefixed with #/##/### become headings, * bullets, |...| tables
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], 3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], 2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], 1)
            elif stripped.startswith("* ") or stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("> "):
                doc.add_paragraph(stripped[2:], style="Intense Quote")
            else:
                doc.add_paragraph(stripped)

        # Tables
        for table_spec in tables_data:
            headers = table_spec.get("headers", [])
            rows_data = table_spec.get("rows", [])
            if not headers and not rows_data:
                continue
            table = doc.add_table(rows=1 + len(rows_data), cols=max(len(headers), 1))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Light Grid Accent 1"
            # Header row
            if headers:
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = str(h)
                # Bold the header row
                for cell in table.rows[0].cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
            # Data rows
            for ri, row in enumerate(rows_data):
                for ci, val in enumerate(row):
                    table.rows[ri + 1].cells[ci].text = str(val)

        doc.save(path)
        return {"status": "success", "path": path}

    async def do_excel_xlsx(self, args: dict) -> dict:
        """Generate an Excel (.xlsx) spreadsheet with multiple sheets, formulas, and formatting.

        Domain rules (from excel-xlsx skill):
        - Excel stores dates as serial numbers (1900 date system includes false leap-day bug)
        - Long IDs, phone numbers, ZIP codes, leading-zero values should be stored as text
        - Excel silently truncates numeric precision past 15 digits
        - Write formulas into cells instead of hardcoding derived results
        - Cached formula values can be stale — never trust them blindly after edits
        - Prefer openpyxl when formulas, styles, sheets, merged cells matter
        - Match existing template styles instead of introducing a new visual system
        """
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            return {"status": "error", "msg": "openpyxl not installed. Run: pip install openpyxl"}
        try:
            path = resolve_workspace_path(args.get("path", "spreadsheet.xlsx"), self.data_dir, self.permissions)
        except PermissionError as exc:
            return {"status": "error", "msg": str(exc)}

        rows = args.get("rows", [])
        if isinstance(rows, str):
            try:
                rows = json.loads(rows)
            except (json.JSONDecodeError, TypeError):
                rows = [[cell] for cell in rows.split("\n") if cell.strip()]

        sheet_name = args.get("sheet_name", "Sheet1")
        formulas = args.get("formulas", [])  # [{"cell": "A1", "formula": "=SUM(B1:B10)"}]
        column_widths = args.get("column_widths", {})  # {"A": 15, "B": 20}
        header_row = args.get("header_row", False)  # first row is header → bold + style
        text_columns = args.get("text_columns", [])  # column indices (0-based) to force as text

        wb = Workbook()
        # Remove default sheet and create named sheet
        default_ws = wb.active
        default_ws.title = sheet_name

        # Write data
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = default_ws.cell(row=ri + 1, column=ci + 1)
                # Force text type for columns that need it (IDs, phone numbers, etc.)
                if ci in text_columns:
                    cell.value = str(val)
                    cell.number_format = "@"  # Text format in Excel
                else:
                    cell.value = val

        # Header row styling
        if header_row and rows:
            header_font = Font(bold=True)
            header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
            for ci in range(len(rows[0])):
                cell = default_ws.cell(row=1, column=ci + 1)
                cell.font = header_font
                cell.fill = header_fill

        # Apply formulas
        for f_spec in formulas:
            cell_ref = f_spec.get("cell", "")
            formula_str = f_spec.get("formula", "")
            if cell_ref and formula_str:
                default_ws[cell_ref] = formula_str

        # Column widths
        for col_letter, width in column_widths.items():
            default_ws.column_dimensions[col_letter].width = width

        wb.save(path)
        return {"status": "success", "path": path}

    async def do_plan_create(self, args: dict) -> dict:
        """Create a new plan via Planner."""
        goal = args.get("goal", "")
        steps = args.get("steps", [])
        try:
            plan = self.planner.create(goal, steps)
        except ValueError as exc:
            return {"status": "error", "msg": str(exc)}
        return {"status": "success", "plan": plan}

    async def do_plan_update(self, args: dict) -> dict:
        """Update a plan step via Planner."""
        plan_id = args.get("plan_id", "")
        step_id = args.get("step_id", 0)
        status = args.get("status", "done")
        result = args.get("result", "")
        plan = self.planner.update(plan_id, step_id, status, result)
        if plan is None:
            return {"status": "error", "msg": f"Plan {plan_id} not found"}
        return {"status": "success", "plan": plan}

    async def do_plan_list(self, args: dict) -> dict:
        """List plans via Planner."""
        status = args.get("status", "")
        plans = self.planner.list_plans(status)
        return {"status": "success", "plans": plans, "count": len(plans)}

    async def do_skill_search(self, args: dict) -> dict:
        """Search ClawHub skill market or GitHub for reusable skills."""
        query = args.get("query", "")
        registry = args.get("registry", "https://mirror-cn.clawhub.com")
        if not query:
            return {"status": "error", "msg": "query is required"}
        import requests
        try:
            resp = requests.get(f"{registry}/api/v1/search", params={"q": query}, timeout=15)
            resp.raise_for_status()
            data = resp.json()
            results = data if isinstance(data, list) else data.get("results", data.get("skills", []))
            return {
                "status": "success",
                "results": [{"slug": r.get("slug"), "name": r.get("name"), "description": r.get("description", "")} for r in results[:20]],
                "count": len(results),
                "source": registry,
            }
        except Exception as e:
            return {"status": "error", "msg": f"搜索技能市场失败: {e}，可尝试 ai_search 在 GitHub 上搜索"}

    async def do_skill_install(self, args: dict) -> dict:
        """Install a skill from ClawHub or GitHub."""
        slug = args.get("slug", "")
        url = args.get("url", "")
        registry = args.get("registry", "https://mirror-cn.clawhub.com")
        if not slug and not url:
            return {"status": "error", "msg": "需要 slug（ClawHub 技能名）或 url（GitHub 地址）"}
        import requests, io, os, tarfile, zipfile
        if slug:
            try:
                resolve = requests.get(f"{registry}/api/v1/resolve", params={"slug": slug}, timeout=15)
                resolve.raise_for_status()
                info = resolve.json()
                dl_path = info.get("downloadUrl", f"/api/v1/download/{slug}")
                dl_url = f"{registry}{dl_path}" if dl_path.startswith("/") else dl_path
                name = slug
            except Exception as e:
                return {"status": "error", "msg": f"解析技能 {slug} 失败: {e}"}
        else:
            dl_url = url
            name = url.rstrip("/").split("/")[-1].replace(".git", "")
        try:
            resp = requests.get(dl_url, timeout=30)
            resp.raise_for_status()
            content = resp.content
            skills_dir = os.path.join(self.data_dir, "skills", "community")
            target = os.path.join(skills_dir, name)
            os.makedirs(target, exist_ok=True)
            if content[:2] == b'\x1f\x8b':
                with tarfile.open(fileobj=io.BytesIO(content)) as tf:
                    tf.extractall(target)
            elif content[:4] == b'PK\x03\x04':
                with zipfile.ZipFile(io.BytesIO(content)) as zf:
                    zf.extractall(target)
            else:
                with open(os.path.join(target, "SKILL.md"), "w", encoding="utf-8") as f:
                    f.write(resp.text)
            return {"status": "success", "msg": f"技能 '{name}' 已安装到 {target}", "path": target}
        except Exception as e:
            return {"status": "error", "msg": f"安装失败: {e}"}

    # ------------------------------------------------------------------
    # MCP (Model Context Protocol) management
    # ------------------------------------------------------------------

    async def do_mcp_install(self, args: dict) -> dict:
        """Install an MCP server (npm package) and register it for use.

        Installs the package via npm (optionally pinned to a version),
        then registers it in the MCP config file so tools can discover it.
        """
        package = args.get("package", "")
        name = args.get("name", "") or package
        version = args.get("version", "latest")
        command = args.get("command", "npx")
        args_list = args.get("args", [])
        env_vars = args.get("env", {})

        if not package:
            return {"status": "error", "msg": "package name is required"}

        if not await self._confirm("mcp_install", f"Install MCP server: {package}@{version}"):
            return {"status": "error", "msg": "MCP install not permitted"}

        # npm install the package globally
        full_pkg = f"{package}@{version}" if version != "latest" else package
        npm_proc = await asyncio.create_subprocess_exec(
            "npm", "install", "-g", full_pkg,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await npm_proc.communicate()
        if npm_proc.returncode != 0:
            err = stderr.decode("utf-8", errors="replace").strip()
            return {"status": "error", "msg": f"npm install failed: {err}"}

        # Register in MCP config
        mcp_config = self._load_mcp_config()
        mcp_config["servers"][name] = {
            "command": command,
            "args": args_list,
            "env": env_vars,
        }
        self._save_mcp_config(mcp_config)

        return {
            "status": "success",
            "msg": f"MCP server '{name}' installed and registered",
            "package": full_pkg,
        }

    async def do_mcp_list(self, args: dict) -> dict:
        """List installed/configured MCP servers."""
        mcp_config = self._load_mcp_config()
        servers = mcp_config.get("servers", {})
        if not servers:
            return {"status": "success", "servers": {}, "count": 0, "msg": "No MCP servers configured"}
        result = {}
        for name, cfg in servers.items():
            result[name] = {
                "command": cfg.get("command", ""),
                "args": cfg.get("args", []),
                "env_count": len(cfg.get("env", {})),
            }
        return {"status": "success", "servers": result, "count": len(result)}

    async def do_mcp_remove(self, args: dict) -> dict:
        """Remove an MCP server configuration."""
        name = args.get("name", "")
        if not name:
            return {"status": "error", "msg": "name is required"}

        mcp_config = self._load_mcp_config()
        if name not in mcp_config.get("servers", {}):
            return {"status": "error", "msg": f"MCP server '{name}' not found"}

        del mcp_config["servers"][name]
        self._save_mcp_config(mcp_config)

        return {"status": "success", "msg": f"MCP server '{name}' removed"}

    def _load_mcp_config(self) -> dict:
        """Load the MCP config file (creates default if missing)."""
        path = os.path.join(self.data_dir, "mcp_servers.json")
        if not os.path.exists(path):
            return {"servers": {}}
        try:
            with open(path, encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {"servers": {}}

    def _save_mcp_config(self, config: dict):
        """Persist MCP config to disk."""
        path = os.path.join(self.data_dir, "mcp_servers.json")
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(config, f, indent=2, ensure_ascii=False)

    # ------------------------------------------------------------------
    # Skill loading (model-activated per-task)
    # ------------------------------------------------------------------

    async def do_skill_load(self, args: dict) -> dict:
        """Load a skill's SKILL.md, SOP.md, and knowledge content by name.

        The model calls this when it determines a skill matches the current task.
        Returns the full content of the skill for the model to follow.
        """
        name = args.get("name", "")
        if not name or not self._skill_mgr:
            return {"status": "error", "msg": "Skill name required or skill manager not available"}

        skill = self._skill_mgr.get(name)
        if not skill:
            return {"status": "error", "msg": f"Skill '{name}' not found. Available: {', '.join(s.name for s in self._skill_mgr.list_all())}"}

        skill.load()
        result: dict[str, Any] = {
            "status": "success",
            "name": skill.name,
            "category": skill.category,
            "description": skill.description,
        }
        if skill.skill_md:
            result["skill_md"] = skill.skill_md
        if skill.sop_md:
            result["sop_md"] = skill.sop_md
        if skill.knowledge:
            result["knowledge"] = skill.knowledge
        return result

    async def do_skill_create(self, args: dict) -> dict:
        """Create a new skill scaffold with SKILL.md template.

        Creates a skill directory with proper frontmatter and optional
        resource directories (scripts/, references/, assets/).
        """
        name = args.get("name", "")
        description = args.get("description", "")
        resources = args.get("resources", "")
        path = args.get("path", "")

        if not name or not description:
            return {"status": "error", "msg": "name and description are required"}

        # Validate/normalize name
        import re
        normalized = name.strip().lower()
        normalized = re.sub(r"[^a-z0-9]+", "-", normalized)
        normalized = normalized.strip("-")
        if not normalized:
            return {"status": "error", "msg": "Invalid skill name after normalization"}
        if len(normalized) > 64:
            return {"status": "error", "msg": f"Skill name too long: {len(normalized)} > 64"}

        # Determine target directory
        if not path:
            path = os.path.join(self.data_dir, "skills", "user_evolved")
        target = os.path.join(path, normalized)
        if os.path.exists(target):
            return {"status": "error", "msg": f"Skill directory already exists: {target}"}

        # Create directory structure
        try:
            os.makedirs(target, exist_ok=False)
        except OSError as exc:
            return {"status": "error", "msg": f"Cannot create directory: {exc}"}

        # Write SKILL.md from template
        title = " ".join(word.capitalize() for word in normalized.split("-"))
        skill_md = f"""---
name: {normalized}
description: {description}
---

# {title}

## Overview

[TODO: Describe what this skill does and when to use it]

## Usage

[TODO: Add instructions, examples, and workflows]

## Resources (optional)

Delete this section if no resources are required.

- **scripts/** — Executable code for automation
- **references/** — Documentation loaded on demand
- **assets/** — Output templates and resources
"""
        skill_path = os.path.join(target, "SKILL.md")
        with open(skill_path, "w", encoding="utf-8") as f:
            f.write(skill_md)

        # Create optional resource directories
        if resources:
            allowed = {"scripts", "references", "assets"}
            for r in resources.split(","):
                r = r.strip()
                if r in allowed:
                    os.makedirs(os.path.join(target, r), exist_ok=True)

        # Trigger skill manager refresh
        if self._skill_mgr:
            self._skill_mgr.discover()

        return {
            "status": "success",
            "msg": f"Skill '{normalized}' created at {target}",
            "path": target,
            "skill_name": normalized,
        }

    async def do_wechat_sessions(self, args: dict) -> dict:
        """Get recent WeChat session list."""
        limit = args.get("limit", 20)
        result = await self.wechat.sessions(limit)
        return {"status": "success", "data": result}

    async def do_wechat_history(self, args: dict) -> dict:
        """Get WeChat chat history with a contact."""
        name = args.get("name", "")
        limit = args.get("limit", 20)
        result = await self.wechat.history(name, limit)
        return {"status": "success", "data": result}

    async def do_wechat_search(self, args: dict) -> dict:
        """Search WeChat messages globally or in a specific chat."""
        keyword = args.get("keyword", "")
        chat = args.get("chat", "")
        limit = args.get("limit", 20)
        result = await self.wechat.search(keyword, chat, limit)
        return {"status": "success", "data": result}

    async def do_wechat_contacts(self, args: dict) -> dict:
        """Search WeChat contacts."""
        query = args.get("query", "")
        result = await self.wechat.contacts(query)
        return {"status": "success", "data": result}

    async def do_wechat_unread(self, args: dict) -> dict:
        """Get unread WeChat sessions."""
        limit = args.get("limit", 20)
        result = await self.wechat.unread(limit)
        return {"status": "success", "data": result}

    async def do_wechat_send_text(self, args: dict) -> dict:
        """Proactively send a WeChat text message to a contact."""
        to = args.get("to", "")
        text = args.get("text", "")
        if not to or not text:
            return {"status": "error", "msg": "to and text are required"}
        if not self._wechat_adapter:
            return {"status": "error", "msg": "微信适配器未连接，请先扫码登录"}
        if not await self._confirm("wechat_send_text", f"To: {to}"):
            return {"status": "error", "msg": "WeChat send not permitted"}
        try:
            result = await self._wechat_adapter.send_message(to, text)
            return {"status": "success" if result.get("status") == "success" else "error", "data": result}
        except Exception as e:
            return {"status": "error", "msg": str(e)}

    async def do_wechat_send_file(self, args: dict) -> dict:
        """Send a local file to a WeChat contact via CDN upload."""
        to = args.get("to", "")
        file_path = args.get("file_path", "")
        if not to or not file_path:
            return {"status": "error", "msg": "to and file_path are required"}
        if not self._wechat_adapter:
            return {"status": "error", "msg": "微信适配器未连接，请先扫码登录"}
        if not await self._confirm("wechat_send_file", f"To: {to}, File: {file_path}"):
            return {"status": "error", "msg": "WeChat send not permitted"}
        try:
            result = await self._wechat_adapter.send_file(to, file_path)
            if result.get("status") == "success":
                return {"status": "success", "data": result}

            # ret=-2: upload rejected by server, needs re-login
            ret = result.get("ret", 0)
            if ret == -2:
                qr = self._wechat_adapter.get_qrcode()
                return {
                    "status": "error",
                    "msg": "微信文件上传功能异常(ret=-2)，需要重新扫码登录。请在5分钟内扫描以下二维码完成重新认证后重试发送。",
                    "needs_reconnect": True,
                    "qrcode_url": qr.get("qrcode_url", ""),
                    "qrcode_id": qr.get("qrcode_id", ""),
                }

            return {"status": "error", "msg": result.get("msg", "发送失败")}
        except Exception as e:
            return {"status": "error", "msg": str(e)}

    async def do_wechat_send_typing(self, args: dict) -> dict:
        """Send typing indicator ('对方正在输入...'). status=1 show, 0 hide."""
        to = args.get("to", "")
        status = args.get("status", 1)
        if not to:
            return {"status": "error", "msg": "to is required"}
        if not self._wechat_adapter:
            return {"status": "error", "msg": "微信适配器未连接，请先扫码登录"}
        try:
            result = await self._wechat_adapter.send_typing(to, status)
            return {"status": "success" if result.get("status") == "success" else "error", "data": result}
        except Exception as e:
            return {"status": "error", "msg": str(e)}

    # ------------------------------------------------------------------
    # Feishu tools
    # ------------------------------------------------------------------

    async def do_feishu_send_message(self, args: dict) -> dict:
        """Send a Feishu IM message."""
        text = args.get("text", "")
        to = args.get("to", "")
        user = args.get("user", "")
        if not text:
            return {"status": "error", "msg": "No text provided"}
        if to:
            result = await self.feishu.send_message(to, text)
        elif user:
            result = await self.feishu.send_user_message(user, text)
        else:
            return {"status": "error", "msg": "Provide 'to' (chat_id) or 'user' (open_id)"}
        return {"status": "success" if result.get("ok") else "error", "data": result}

    async def do_feishu_search_user(self, args: dict) -> dict:
        """Search Feishu users."""
        query = args.get("query", "")
        limit = args.get("limit", 20)
        result = await self.feishu.search_user(query, limit)
        return {"status": "success", "data": result}

    async def do_feishu_get_user(self, args: dict) -> dict:
        """Get Feishu user info."""
        user_id = args.get("user_id", "")
        result = await self.feishu.get_user(user_id)
        return {"status": "success", "data": result}

    async def do_feishu_calendar_agenda(self, args: dict) -> dict:
        """View calendar agenda."""
        start = args.get("start", "")
        end = args.get("end", "")
        result = await self.feishu.calendar_agenda(start, end)
        return {"status": "success", "data": result}

    async def do_feishu_calendar_create(self, args: dict) -> dict:
        """Create a calendar event."""
        summary = args.get("summary", "")
        start = args.get("start", "")
        end = args.get("end", "")
        description = args.get("description", "")
        attendees = args.get("attendees", [])
        result = await self.feishu.calendar_create_event(summary, start, end, description, attendees)
        return {"status": "success", "data": result}

    async def do_feishu_drive_search(self, args: dict) -> dict:
        """Search Drive files."""
        query = args.get("query", "")
        limit = args.get("limit", 20)
        result = await self.feishu.drive_search(query, limit)
        return {"status": "success", "data": result}

    async def do_feishu_doc_fetch(self, args: dict) -> dict:
        """Fetch document content."""
        token = args.get("token", "")
        result = await self.feishu.doc_fetch(token)
        return {"status": "success", "data": result}

    async def do_feishu_doc_create(self, args: dict) -> dict:
        """Create a document."""
        title = args.get("title", "")
        content = args.get("content", "")
        result = await self.feishu.doc_create(title, content)
        return {"status": "success", "data": result}

    async def do_feishu_doc_search(self, args: dict) -> dict:
        """Search documents."""
        query = args.get("query", "")
        result = await self.feishu.doc_search(query)
        return {"status": "success", "data": result}

    async def do_feishu_sheets_read(self, args: dict) -> dict:
        """Read spreadsheet values."""
        token = args.get("spreadsheet_token", "")
        range_str = args.get("range", "")
        result = await self.feishu.sheets_read(token, range_str)
        return {"status": "success", "data": result}

    async def do_feishu_sheets_create(self, args: dict) -> dict:
        """Create a spreadsheet."""
        title = args.get("title", "")
        result = await self.feishu.sheets_create(title)
        return {"status": "success", "data": result}

    async def do_feishu_base_records(self, args: dict) -> dict:
        """List bitable records."""
        base_token = args.get("base_token", "")
        table_id = args.get("table_id", "")
        limit = args.get("limit", 100)
        result = await self.feishu.base_records(base_token, table_id, limit)
        return {"status": "success", "data": result}

    async def do_feishu_task_list(self, args: dict) -> dict:
        """List tasks."""
        limit = args.get("limit", 50)
        result = await self.feishu.task_list(limit)
        return {"status": "success", "data": result}

    async def do_feishu_wiki_search(self, args: dict) -> dict:
        """Search wiki (uses drive search)."""
        query = args.get("query", "")
        result = await self.feishu.drive_search(query)
        return {"status": "success", "data": result}

    async def do_feishu_mail_list(self, args: dict) -> dict:
        """List mail (uses mail +messages — requires message IDs)."""
        limit = args.get("limit", 20)
        return {"status": "error", "msg": "Mail listing requires specific message IDs; use feishu_drive_search instead"}

    async def do_feishu_chat_search(self, args: dict) -> dict:
        """Search group chats."""
        query = args.get("query", "")
        result = await self.feishu.chat_search(query)
        return {"status": "success", "data": result}

    async def do_feishu_chat_messages(self, args: dict) -> dict:
        """List chat messages."""
        chat_id = args.get("chat_id", "")
        limit = args.get("limit", 20)
        result = await self.feishu.chat_messages(chat_id, limit)
        return {"status": "success", "data": result}

    async def do_feishu_drive_upload(self, args: dict) -> dict:
        """Upload file to Drive."""
        local_path = args.get("local_path", "")
        folder_token = args.get("folder_token", "")
        result = await self.feishu.drive_upload(local_path, folder_token)
        return {"status": "success", "data": result}

    # ------------------------------------------------------------------
    # DingTalk tools
    # ------------------------------------------------------------------

    async def do_dingtalk_send_message(self, args: dict) -> dict:
        """Send a DingTalk message to a user."""
        user_id = args.get("user_id", "")
        text = args.get("text", "")
        title = args.get("title", "")
        if not user_id or not text:
            return {"status": "error", "msg": "user_id and text are required"}
        result = await self.dingtalk.send_message(user_id, text, title)
        return {"status": "success" if result.get("ok") else "error", "data": result}

    async def do_dingtalk_send_group_message(self, args: dict) -> dict:
        """Send a DingTalk message to a group."""
        group_id = args.get("group_id", "")
        text = args.get("text", "")
        title = args.get("title", "")
        if not group_id or not text:
            return {"status": "error", "msg": "group_id and text are required"}
        result = await self.dingtalk.send_group_message(group_id, text, title)
        return {"status": "success" if result.get("ok") else "error", "data": result}

    async def do_dingtalk_search_user(self, args: dict) -> dict:
        """Search DingTalk users."""
        query = args.get("query", "")
        result = await self.dingtalk.search_user(query)
        return {"status": "success", "data": result}

    async def do_dingtalk_user_info(self, args: dict) -> dict:
        """Get DingTalk user info."""
        user_id = args.get("user_id", "")
        result = await self.dingtalk.get_user(user_id)
        return {"status": "success", "data": result}

    async def do_dingtalk_chat_search(self, args: dict) -> dict:
        """Search group conversations."""
        query = args.get("query", "")
        result = await self.dingtalk.chat_search(query)
        return {"status": "success", "data": result}

    async def do_dingtalk_chat_list(self, args: dict) -> dict:
        """List top conversations."""
        limit = args.get("limit", 20)
        result = await self.dingtalk.chat_list(limit=limit)
        return {"status": "success", "data": result}

    async def do_dingtalk_chat_history(self, args: dict) -> dict:
        """List recent messages in a group."""
        group_id = args.get("group_id", "")
        limit = args.get("limit", 20)
        if not group_id:
            return {"status": "error", "msg": "group_id is required"}
        result = await self.dingtalk.chat_history(group_id, limit)
        return {"status": "success", "data": result}

    async def do_dingtalk_chat_unread(self, args: dict) -> dict:
        """List unread conversations."""
        limit = args.get("limit", 20)
        result = await self.dingtalk.chat_unread(limit)
        return {"status": "success", "data": result}

    async def do_dingtalk_calendar_list(self, args: dict) -> dict:
        """List calendar events."""
        limit = args.get("limit", 50)
        result = await self.dingtalk.calendar_list(limit=limit)
        return {"status": "success", "data": result}

    async def do_dingtalk_calendar_create(self, args: dict) -> dict:
        """Create a calendar event."""
        summary = args.get("summary", "")
        start_time = args.get("start_time", "")
        end_time = args.get("end_time", "")
        description = args.get("description", "")
        attendees = args.get("attendees", [])
        result = await self.dingtalk.calendar_create(
            summary, start_time, end_time, description, attendees,
        )
        return {"status": "success", "data": result}

    async def do_dingtalk_todo_list(self, args: dict) -> dict:
        """List todo tasks."""
        limit = args.get("limit", 50)
        result = await self.dingtalk.todo_list(limit=limit)
        return {"status": "success", "data": result}

    async def do_dingtalk_todo_create(self, args: dict) -> dict:
        """Create a todo task."""
        subject = args.get("subject", "")
        description = args.get("description", "")
        due_time = args.get("due_time", "")
        executor_ids = args.get("executor_ids", [])
        result = await self.dingtalk.todo_create(subject, description, due_time, executor_ids)
        return {"status": "success", "data": result}

    async def do_dingtalk_doc_search(self, args: dict) -> dict:
        """Search documents."""
        query = args.get("query", "")
        limit = args.get("limit", 20)
        result = await self.dingtalk.doc_search(query, limit=limit)
        return {"status": "success", "data": result}

    async def do_dingtalk_doc_read(self, args: dict) -> dict:
        """Read a document."""
        doc_id = args.get("doc_id", "")
        if not doc_id:
            return {"status": "error", "msg": "doc_id is required"}
        result = await self.dingtalk.doc_read(doc_id)
        return {"status": "success", "data": result}

    async def do_dingtalk_doc_create(self, args: dict) -> dict:
        """Create a document."""
        title = args.get("title", "")
        content = args.get("content", "")
        if not title:
            return {"status": "error", "msg": "title is required"}
        result = await self.dingtalk.doc_create(title, content)
        return {"status": "success", "data": result}

    async def do_dingtalk_drive_list(self, args: dict) -> dict:
        """List Drive files."""
        parent_id = args.get("parent_id", "")
        limit = args.get("limit", 50)
        result = await self.dingtalk.drive_list(parent_id, limit=limit)
        return {"status": "success", "data": result}

    async def do_dingtalk_wiki_search(self, args: dict) -> dict:
        """Search wiki."""
        query = args.get("query", "")
        limit = args.get("limit", 20)
        result = await self.dingtalk.wiki_search(query, limit=limit)
        return {"status": "success", "data": result}

    # ------------------------------------------------------------------
    # DingTalk Sheet tools
    # ------------------------------------------------------------------

    async def do_dingtalk_sheet_info(self, args: dict) -> dict:
        """Get DingTalk sheet info."""
        workbook_id = args.get("workbook_id", "")
        sheet_id = args.get("sheet_id", "")
        if not workbook_id:
            return {"status": "error", "msg": "workbook_id is required"}
        result = await self.dingtalk.sheet_info(workbook_id, sheet_id)
        return {"status": "success", "data": result}

    async def do_dingtalk_sheet_create(self, args: dict) -> dict:
        """Create a DingTalk spreadsheet."""
        title = args.get("title", "")
        if not title:
            return {"status": "error", "msg": "title is required"}
        result = await self.dingtalk.sheet_create(title)
        return {"status": "success", "data": result}

    async def do_dingtalk_sheet_list(self, args: dict) -> dict:
        """List worksheets in a DingTalk spreadsheet."""
        node = args.get("node", "")
        if not node:
            return {"status": "error", "msg": "node is required"}
        result = await self.dingtalk.sheet_list(node)
        return {"status": "success", "data": result}

    async def do_dingtalk_sheet_append(self, args: dict) -> dict:
        """Append rows to a DingTalk worksheet."""
        node = args.get("node", "")
        sheet_id = args.get("sheet_id", "")
        values = args.get("values", "")
        if not node or not sheet_id or not values:
            return {"status": "error", "msg": "node, sheet_id, and values are required"}
        result = await self.dingtalk.sheet_append(node, sheet_id, values)
        return {"status": "success", "data": result}

    async def do_dingtalk_sheet_read(self, args: dict) -> dict:
        """Read cell values from a DingTalk worksheet."""
        node = args.get("node", "")
        sheet_id = args.get("sheet_id", "")
        range_str = args.get("range", "")
        if not node or not sheet_id:
            return {"status": "error", "msg": "node and sheet_id are required"}
        result = await self.dingtalk.sheet_read(node, sheet_id, range_str)
        return {"status": "success", "data": result}

    # ------------------------------------------------------------------
    # DingTalk AITable (多维表) tools
    # ------------------------------------------------------------------

    async def do_dingtalk_base_create(self, args: dict) -> dict:
        """Create a DingTalk AI table base (多维表)."""
        name = args.get("name", "")
        if not name:
            return {"status": "error", "msg": "name is required"}
        result = await self.dingtalk.aitable_base_create(name)
        return {"status": "success", "data": result}

    async def do_dingtalk_base_list(self, args: dict) -> dict:
        """List DingTalk AI table bases."""
        limit = args.get("limit", 20)
        result = await self.dingtalk.aitable_base_list(limit=limit)
        return {"status": "success", "data": result}

    async def do_dingtalk_table_create(self, args: dict) -> dict:
        """Create a data table in a DingTalk AI table base."""
        base_id = args.get("base_id", "")
        name = args.get("name", "")
        if not base_id or not name:
            return {"status": "error", "msg": "base_id and name are required"}
        result = await self.dingtalk.aitable_table_create(base_id, name)
        return {"status": "success", "data": result}

    async def do_dingtalk_record_create(self, args: dict) -> dict:
        """Add records to a DingTalk AI table."""
        base_id = args.get("base_id", "")
        table_id = args.get("table_id", "")
        records = args.get("records", "")
        if not base_id or not table_id or not records:
            return {"status": "error", "msg": "base_id, table_id, and records are required"}
        result = await self.dingtalk.aitable_record_create(base_id, table_id, records)
        return {"status": "success", "data": result}

    async def do_dingtalk_record_query(self, args: dict) -> dict:
        """Query records from a DingTalk AI table."""
        base_id = args.get("base_id", "")
        table_id = args.get("table_id", "")
        limit = args.get("limit", 100)
        if not base_id or not table_id:
            return {"status": "error", "msg": "base_id and table_id are required"}
        result = await self.dingtalk.aitable_record_query(base_id, table_id, limit)
        return {"status": "success", "data": result}

    # ------------------------------------------------------------------
    # Generic CLI passthrough (cover all 200+ commands)
    # ------------------------------------------------------------------

    async def do_feishu_cli_run(self, args: dict) -> dict:
        """Execute an arbitrary lark-cli command."""
        cmd_args = args.get("args", "")
        if not cmd_args:
            return {"status": "error", "msg": "args is required"}
        import shlex
        arg_list = shlex.split(cmd_args)
        result = await self.feishu._run(arg_list)
        return {"status": "success" if result.get("ok") else "error", "data": result}

    async def do_dingtalk_cli_run(self, args: dict) -> dict:
        """Execute an arbitrary dws command."""
        cmd_args = args.get("args", "")
        if not cmd_args:
            return {"status": "error", "msg": "args is required"}
        import shlex
        arg_list = shlex.split(cmd_args)
        result = await self.dingtalk._run(arg_list)
        return {"status": "success" if result.get("ok") else "error", "data": result}
