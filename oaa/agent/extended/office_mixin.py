"""Office mixin — Word and Excel document generation tools."""
import json
import os
from ..tool_decorator import agent_tool
from ..path_utils import resolve_workspace_path


class OfficeMixin:
    """Word and Excel document generation tools."""

    async def do_word_doc(self, args: dict) -> dict:
        """Generate a Word (.docx) document with headings, tables, paragraphs, and styles.

        Domain rules (from word-docx skill):
        - Prefer named styles (Heading 1/2/3, Normal) over direct formatting
        - A .docx is a ZIP of XML parts — structure matters as much as visible text
        - Margins, orientation, headers/footers are section-level, not document-level
        - Lists and numbering use Word's numbering definitions, not Unicode bullets
        - Tracked changes, comments, and fields need precise edits — visible text is
          not the full document when revisions are enabled
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            return {"status": "error", "msg": "python-docx not installed. Run: pip install python-docx"}
        try:
            path = resolve_workspace_path(args.get("path", "document.docx"), self.data_dir, self.permissions)
        except PermissionError as exc:
            return {"status": "error", "msg": str(exc)}

        title = args.get("title", "Document")
        content = args.get("content", "")
        tables_data = args.get("tables", [])
        page_orientation = args.get("page_orientation", "portrait")  # 'portrait' or 'landscape'
        margins = args.get("margins", None)  # dict with top/bottom/left/right in inches

        doc = Document()

        # Page setup
        section = doc.sections[0]
        if page_orientation == "landscape":
            section.orientation = 1  # WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        if margins:
            margin_inches = {k: Inches(v) for k, v in margins.items() if k in ("top", "bottom", "left", "right")}
            for k, v in margin_inches.items():
                setattr(section, k, v)

        # Title
        doc.add_heading(title, 0)

        # Parse content: lines prefixed with #/##/### become headings, * bullets, |...| tables
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], 3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], 2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], 1)
            elif stripped.startswith("* ") or stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("> "):
                doc.add_paragraph(stripped[2:], style="Intense Quote")
            else:
                doc.add_paragraph(stripped)

        # Tables
        for table_spec in tables_data:
            headers = table_spec.get("headers", [])
            rows_data = table_spec.get("rows", [])
            if not headers and not rows_data:
                continue
            table = doc.add_table(rows=1 + len(rows_data), cols=max(len(headers), 1))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Light Grid Accent 1"
            # Header row
            if headers:
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = str(h)
                # Bold the header row
                for cell in table.rows[0].cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
            # Data rows
            for ri, row in enumerate(rows_data):
                for ci, val in enumerate(row):
                    table.rows[ri + 1].cells[ci].text = str(val)

        doc.save(path)
        return {"status": "success", "path": path}

    async def do_excel_xlsx(self, args: dict) -> dict:
        """Generate an Excel (.xlsx) spreadsheet with multiple sheets, formulas, and formatting.

        Domain rules (from excel-xlsx skill):
        - Excel stores dates as serial numbers (1900 date system includes false leap-day bug)
        - Long IDs, phone numbers, ZIP codes, leading-zero values should be stored as text
        - Excel silently truncates numeric precision past 15 digits
        - Write formulas into cells instead of hardcoding derived results
        - Cached formula values can be stale — never trust them blindly after edits
        - Prefer openpyxl when formulas, styles, sheets, merged cells matter
        - Match existing template styles instead of introducing a new visual system
        """
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            return {"status": "error", "msg": "openpyxl not installed. Run: pip install openpyxl"}
        try:
            path = resolve_workspace_path(args.get("path", "spreadsheet.xlsx"), self.data_dir, self.permissions)
        except PermissionError as exc:
            return {"status": "error", "msg": str(exc)}

        rows = args.get("rows", [])
        if isinstance(rows, str):
            try:
                rows = json.loads(rows)
            except (json.JSONDecodeError, TypeError):
                rows = [[cell] for cell in rows.split("\n") if cell.strip()]

        sheet_name = args.get("sheet_name", "Sheet1")
        formulas = args.get("formulas", [])  # [{"cell": "A1", "formula": "=SUM(B1:B10)"}]
        column_widths = args.get("column_widths", {})  # {"A": 15, "B": 20}
        header_row = args.get("header_row", False)  # first row is header → bold + style
        text_columns = args.get("text_columns", [])  # column indices (0-based) to force as text

        wb = Workbook()
        # Remove default sheet and create named sheet
        default_ws = wb.active
        default_ws.title = sheet_name

        # Write data
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = default_ws.cell(row=ri + 1, column=ci + 1)
                # Force text type for columns that need it (IDs, phone numbers, etc.)
                if ci in text_columns:
                    cell.value = str(val)
                    cell.number_format = "@"  # Text format in Excel
                else:
                    cell.value = val

        # Header row styling
        if header_row and rows:
            header_font = Font(bold=True)
            header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
            for ci in range(len(rows[0])):
                cell = default_ws.cell(row=1, column=ci + 1)
                cell.font = header_font
                cell.fill = header_fill

        # Apply formulas
        for f_spec in formulas:
            cell_ref = f_spec.get("cell", "")
            formula_str = f_spec.get("formula", "")
            if cell_ref and formula_str:
                default_ws[cell_ref] = formula_str

        # Column widths
        for col_letter, width in column_widths.items():
            default_ws.column_dimensions[col_letter].width = width

        wb.save(path)
        return {"status": "success", "path": path}
